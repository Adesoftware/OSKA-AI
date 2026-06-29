from docx import Document
from reportlab.pdfgen import canvas

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
import threading
import time
import os
import sys
def resource_path(relative_path):
    """
    Get absolute path to resource.
    Works in development and when packaged with PyInstaller.
    """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)
from ollama_client import ask_oska
# ----------------------------
# Main Window
# ----------------------------
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox
from PIL import Image, ImageTk
import threading
import time
import requests
from ollama_client import ask_oska
# -----------------------------
# Splash Screen
# -----------------------------
splash = tk.Tk()
splash.overrideredirect(True)
splash.geometry("500x300")

# Center the splash screen
screen_width = splash.winfo_screenwidth()
screen_height = splash.winfo_screenheight()

x = (screen_width // 2) - (500 // 2)
y = (screen_height // 2) - (300 // 2)

splash.geometry(f"500x300+{x}+{y}")
splash.configure(bg="#0B4F9C")

# Logo
try:
    logo = tk.PhotoImage(
    file=resource_path("assets/OSKA-AI-LOGO.png")
)
    logo = logo.resize((90,90))
    logo_img = ImageTk.PhotoImage(logo)

    tk.Label(
        splash,
        image=logo_img,
        bg="#0B4F9C"
    ).pack(pady=(25,10))
except:
    pass

tk.Label(
    splash,
    text="OSKA AI",
    font=("Segoe UI",28,"bold"),
    fg="white",
    bg="#0B4F9C"
).pack()

tk.Label(
    splash,
    text="Enterprise Intelligence Without the Cloud",
    font=("Segoe UI",11),
    fg="white",
    bg="#0B4F9C"
).pack()

loading = tk.Label(
    splash,
    text="Loading AI Engine...",
    font=("Segoe UI",10),
    fg="white",
    bg="#0B4F9C"
)

loading.pack(pady=30)

splash.update()

time.sleep(7)

splash.destroy()
root = tk.Tk()
root.title("OSKA AI")
logo = tk.PhotoImage(
    file=resource_path("assets/OSKA-AI-LOGO.png")
)

root.iconphoto(True, logo)
root.iconphoto(True, logo)
root.geometry("950x700")
root.configure(bg="#F3F6FB")
root.resizable(True, True)
# ==========================
# Menu Functions
# ==========================

def clear_all():
    input_box.delete("1.0", tk.END)
    output_box.delete("1.0", tk.END)
    status.config(text="Ready", fg="green")

def copy_response():
    text = output_box.get("1.0", tk.END)
    root.clipboard_clear()
    root.clipboard_append(text)
    status.config(text="✓ Response copied", fg="green")
def save_response():

    text = output_box.get("1.0", tk.END).strip()

    if not text:
        messagebox.showwarning(
            "Nothing to Save",
            "There is no response to save."
        )
        return
def export_word():

    text = output_box.get("1.0", tk.END).strip()

    if not text:
        messagebox.showwarning(
            "Nothing to Export",
            "There is no response to export."
        )
        return

    filename = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        title="Export to Word"
    )

    if filename:

        doc = Document()

        doc.add_heading("OSKA AI Response", level=1)

        doc.add_paragraph(text)

        doc.save(filename)

        status.config(
            text="✓ Word document exported successfully",
            fg="green"
        )
        
        
    filename = filedialog.asksaveasfilename(
        defaultextension=".txt",
        filetypes=[
            ("Text Files", "*.txt"),
            ("All Files", "*.*")
        ],
        title="Save OSKA Response"
    )
def export_pdf():

    text = output_box.get("1.0", tk.END).strip()

    if not text:
        messagebox.showwarning(
            "Nothing to Export",
            "There is no response to export."
        )
        return

    filename = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF Files", "*.pdf")],
        title="Export to PDF"
    )

    if filename:

        pdf = canvas.Canvas(filename)

        pdf.setTitle("OSKA AI Response")

        y = 800

        for line in text.split("\n"):

            pdf.drawString(40, y, line)

            y -= 18

            if y < 40:

                pdf.showPage()

                y = 800

        pdf.save()

        status.config(
            text="✓ PDF exported successfully",
            fg="green"
        )
    if filename:

        with open(filename, "w", encoding="utf-8") as file:
            file.write(text)

        status.config(
            text="✓ Response saved successfully",
            fg="green"
        )
def show_about():
    messagebox.showinfo(
        "About OSKA AI",
        """OSKA AI v1.0

Smart Knowledge Assistant

Enterprise Intelligence Without the Cloud

Developer:
Adeola Abolarin ADEOLAABORIN2@gmail.com

Built With:
• Python
• Tkinter
• Ollama
• Qwen2.5:1.5B

© 2026 All Rights Reserved
"""
    )

# ==========================
# Menu Bar
# ==========================

menu_bar = tk.Menu(root)

# File Menu
file_menu = tk.Menu(menu_bar, tearoff=0)

file_menu.add_command(label="New", command=clear_all)

file_menu.add_command(label="Save as Text", command=save_response)

file_menu.add_command(label="Export to Word", command=export_word)

file_menu.add_command(label="Export to PDF", command=export_pdf)

file_menu.add_separator()

file_menu.add_command(label="Exit", command=root.quit)
menu_bar.add_cascade(label="File", menu=file_menu)
# Edit Menu
edit_menu = tk.Menu(menu_bar, tearoff=0)
edit_menu.add_command(label="Clear", command=clear_all)
edit_menu.add_command(label="Copy Response", command=copy_response)
menu_bar.add_cascade(label="Edit", menu=edit_menu)

# Help Menu
help_menu = tk.Menu(menu_bar, tearoff=0)
help_menu.add_command(label="About OSKA AI", command=show_about)
menu_bar.add_cascade(label="Help", menu=help_menu)

root.config(menu=menu_bar)
# ----------------------------
# Generate Response
# ----------------------------

def generate_response():

    prompt = input_box.get("1.0", tk.END).strip()

    if prompt == "":
        status.config(text="Please enter a request.", fg="red")
        return

    output_box.delete("1.0", tk.END)

    generate_btn.config(state="disabled")

    threading.Thread(
        target=run_ai,
        args=(prompt,),
        daemon=True
    ).start()


def run_ai(prompt):

    start = time.time()


    start = time.time()

    animation = [
        "Analyzing your request.",
        "Analyzing your request..",
        "Analyzing your request..."
    ]

    for i in range(3):
        status.config(
            text=animation[i],
            fg="blue"
        )
        time.sleep(0.4)

    answer = ask_oska(prompt)
    

    answer = ask_oska(prompt)

    elapsed = round(time.time() - start, 1)

    root.after(0, lambda: output_box.delete("1.0", tk.END))
    root.after(0, lambda: output_box.insert(tk.END, answer))
    root.after(0, lambda: status.config(
        text=f"✓ Completed in {elapsed} seconds",
        fg="green"
    ))
    root.after(0, lambda: generate_btn.config(state="normal"))

    for i in range(3):
        status.config(
            text=animation[i],
            fg="blue"
        )
        time.sleep(0.4)

    answer = ask_oska(prompt)

    elapsed = round(time.time() - start,1)

    output_box.insert(tk.END, answer)

    status.config(
        text=f"✓ Completed in {elapsed} seconds",
        fg="green"
    )

    generate_btn.config(state="normal")
# ----------------------------
# Header
# ----------------------------
header = tk.Frame(root, bg="#0B4F9C", height=90)
header.pack(fill="x")

title = tk.Label(
    header,
    text="OSKA AI",
    bg="#0B4F9C",
    fg="white",
    font=("Segoe UI", 24, "bold")
)
title.pack(pady=(15, 0))

subtitle = tk.Label(
    header,
    text="Enterprise Intelligence Without the Cloud",
    bg="#0B4F9C",
    fg="white",
    font=("Segoe UI", 11)
)
subtitle.pack()

# ----------------------------
# Main Container
# ----------------------------
main = tk.Frame(root, bg="#F3F6FB")
main.pack(fill="both", expand=True, padx=20, pady=20)

# Feature
tk.Label(
    main,
    text="Select Feature",
    font=("Segoe UI", 11, "bold"),
    bg="#F3F6FB"
).pack(anchor="w")

feature = ttk.Combobox(
    main,
    values=[
        "Business Letter",
        "Memo Generator",
        "Text Summarizer",
        "General AI Chat"
    ],
    state="readonly",
    width=35
)
feature.current(0)
feature.pack(anchor="w", pady=(5, 15))

# Input
tk.Label(
    main,
    text="Enter your request",
    font=("Segoe UI", 11, "bold"),
    bg="#F3F6FB"
).pack(anchor="w")

input_box = scrolledtext.ScrolledText(
    main,
    height=7,
    font=("Segoe UI", 10)
)
input_box.pack(fill="x", pady=8)

# Buttons
button_frame = tk.Frame(main, bg="#F3F6FB")
button_frame.pack(fill="x", pady=10)

generate_btn = tk.Button(
    button_frame,
    text="Generate",
    bg="#0B4F9C",
    fg="white",
    font=("Segoe UI", 11, "bold"),
    width=15,
    command=generate_response
)
generate_btn.pack(side="left", padx=5)


clear_btn = tk.Button(
    button_frame,
    text="Clear",
    width=12,
    command=lambda: (
        input_box.delete("1.0", tk.END),
        output_box.delete("1.0", tk.END),
        status.config(text="Ready", fg="green")
    )
)

clear_btn.pack(side="left", padx=5)

copy_btn = tk.Button(
    button_frame,
    text="Copy",
    width=12,
    command=lambda: root.clipboard_append(output_box.get("1.0", tk.END))
)
copy_btn.pack(side="left", padx=5)

# Status
status = tk.Label(
    main,
    text="Ready",
    fg="green",
    bg="#F3F6FB",
    font=("Segoe UI", 10)
)
status.pack(anchor="w", pady=5)

# Output
tk.Label(
    main,
    text="OSKA Response",
    font=("Segoe UI", 11, "bold"),
    bg="#F3F6FB"
).pack(anchor="w")

output_box = scrolledtext.ScrolledText(
    main,
    height=15,
    font=("Segoe UI", 10)
)
output_box.pack(fill="both", expand=True, pady=8)

root.mainloop()